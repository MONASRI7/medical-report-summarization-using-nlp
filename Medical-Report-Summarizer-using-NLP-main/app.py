from flask import Flask, render_template, request, jsonify, send_file
import nltk
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
import json
import os
from datetime import datetime
import re
from werkzeug.utils import secure_filename
import pandas as pd
import chardet
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import io
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import sqlite3

app = Flask(__name__)

# Configure upload folder and database
UPLOAD_FOLDER = 'uploads'
DB_PATH = 'instance/history.db'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs('instance', exist_ok=True)

def init_db():
    """Initialize the SQLite database."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            summary TEXT NOT NULL,
            key_info TEXT NOT NULL,
            confidence_scores TEXT NOT NULL,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')

# Medical terminology dictionary (simplified version)
MEDICAL_TERMS = {
    'diagnosis': ['diagnosed', 'diagnosis', 'condition', 'disease', 'syndrome'],
    'medications': ['prescribed', 'medication', 'drug', 'medicine', 'treatment'],
    'procedures': ['procedure', 'surgery', 'operation', 'intervention'],
    'symptoms': ['symptom', 'sign', 'complaint', 'presentation'],
    'tests': ['test', 'exam', 'scan', 'imaging', 'laboratory']
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_text_file(file_path):
    """Read text file with proper encoding detection."""
    try:
        # First try to detect the encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding'] if result['encoding'] else 'utf-8'
        
        # Try to read with detected encoding
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    except UnicodeDecodeError:
        # If the detected encoding fails, try with different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode file with any supported encoding")

def read_pdf_file(file_path):
    """Read PDF file and extract text."""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")

def read_docx_file(file_path):
    """Read DOCX file and extract text."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")

def extract_text_from_file(file_path, file_extension):
    """Extract text from different file types."""
    if file_extension == 'txt':
        return read_text_file(file_path)
    elif file_extension == 'pdf':
        return read_pdf_file(file_path)
    elif file_extension in ['doc', 'docx']:
        return read_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def preprocess_text(text):
    """Preprocess the text for better summarization."""
    # Remove extra whitespace and newlines
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Tokenize into sentences
    sentences = sent_tokenize(text)
    
    # Tokenize words and remove stopwords
    stop_words = set(stopwords.words('english'))
    processed_sentences = []
    
    for sentence in sentences:
        # Clean the sentence
        sentence = re.sub(r'[^\w\s.]', '', sentence)
        words = word_tokenize(sentence.lower())
        words = [word for word in words if word.isalnum() and word not in stop_words]
        if words:  # Only add non-empty sentences
            processed_sentences.append(' '.join(words))
    
    return processed_sentences

def generate_summary(text, num_sentences=3):
    """Generate summary using TF-IDF and sentence scoring."""
    if not text.strip():
        return "No text content found in the document."
    
    # Preprocess the text
    processed_sentences = preprocess_text(text)
    
    if not processed_sentences:
        return "Unable to extract meaningful content from the document."
    
    # Create TF-IDF vectorizer
    vectorizer = TfidfVectorizer()
    try:
        tfidf_matrix = vectorizer.fit_transform(processed_sentences)
        
        # Calculate sentence scores based on TF-IDF
        sentence_scores = np.array(tfidf_matrix.sum(axis=1)).flatten()
        
        # Get top sentences
        top_sentence_indices = sentence_scores.argsort()[-num_sentences:][::-1]
        top_sentence_indices = sorted(top_sentence_indices)
        
        # Get original sentences
        original_sentences = sent_tokenize(text)
        summary = [original_sentences[i] for i in top_sentence_indices]
        
        return ' '.join(summary)
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def extract_key_information(text):
    """Extract key information from the text."""
    sentences = sent_tokenize(text)
    key_info = {
        'diagnosis': [],
        'medications': [],
        'recommendations': [],
        'procedures': [],
        'symptoms': [],
        'tests': []
    }
    
    confidence_scores = {}
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        
        # Check each category
        for category, keywords in MEDICAL_TERMS.items():
            matches = [word for word in keywords if word in sentence_lower]
            if matches:
                # Calculate confidence score based on number of matching keywords
                confidence = min(len(matches) / len(keywords), 1.0)
                
                if category == 'diagnosis':
                    key_info['diagnosis'].append(sentence)
                    confidence_scores['diagnosis'] = confidence
                elif category == 'medications':
                    key_info['medications'].append(sentence)
                    confidence_scores['medications'] = confidence
                elif category == 'procedures':
                    key_info['procedures'].append(sentence)
                    confidence_scores['procedures'] = confidence
                elif category == 'symptoms':
                    key_info['symptoms'].append(sentence)
                    confidence_scores['symptoms'] = confidence
                elif category == 'tests':
                    key_info['tests'].append(sentence)
                    confidence_scores['tests'] = confidence
        
        # Check for recommendations
        if any(word in sentence_lower for word in ['recommend', 'advise', 'suggest']):
            key_info['recommendations'].append(sentence)
            confidence_scores['recommendations'] = 0.8
    
    return key_info, confidence_scores

def highlight_medical_terms(text):
    """Highlight medical terms in the text."""
    highlighted_text = text
    for category, terms in MEDICAL_TERMS.items():
        for term in terms:
            pattern = r'\b' + re.escape(term) + r'\b'
            highlighted_text = re.sub(pattern, f'<span class="highlight-{category}">{term}</span>', 
                                    highlighted_text, flags=re.IGNORECASE)
    return highlighted_text

def load_history():
    """Load history from SQLite database."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('SELECT * FROM history ORDER BY timestamp DESC')
    rows = c.fetchall()
    conn.close()
    
    history = []
    for row in rows:
        history.append({
            'id': row[0],
            'filename': row[1],
            'summary': row[2],
            'key_info': json.loads(row[3]),
            'confidence_scores': json.loads(row[4]),
            'timestamp': row[5]
        })
    return history

def save_history(entry):
    """Save entry to SQLite database."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        INSERT INTO history (filename, summary, key_info, confidence_scores)
        VALUES (?, ?, ?, ?)
    ''', (
        entry['filename'],
        entry['summary'],
        json.dumps(entry['key_info']),
        json.dumps(entry.get('confidence_scores', {}))
    ))
    conn.commit()
    conn.close()

def create_pdf_summary(entry, output_path):
    """Create a PDF summary with proper formatting and spacing."""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=10
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=12,
        leading=14,
        spaceAfter=12
    )

    # Title
    story.append(Paragraph("Medical Report Summary", title_style))
    story.append(Spacer(1, 20))

    # Summary section
    story.append(Paragraph("Summary", heading_style))
    # Process summary text to handle line breaks properly
    summary_text = entry['summary'].replace('\n', '<br/>')
    story.append(Paragraph(summary_text, body_style))
    story.append(Spacer(1, 20))

    # Key Information section
    story.append(Paragraph("Key Information", heading_style))
    
    # Create table data
    data = [['Category', 'Information', 'Confidence']]
    for category, items in entry['key_info'].items():
        confidence = entry['confidence_scores'].get(category, 0)
        confidence_text = f"{confidence*100:.1f}%" if confidence else "N/A"
        # Process items to handle line breaks properly
        items_text = "<br/><br/>".join(item.replace('\n', '<br/>') for item in items)
        data.append([
            Paragraph(category, body_style),
            Paragraph(items_text, body_style),
            Paragraph(confidence_text, body_style)
        ])

    # Create and style the table
    table = Table(data, colWidths=[1.5*inch, 3.5*inch, 1*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))

    story.append(table)
    story.append(Spacer(1, 20))

    # Build the PDF
    doc.build(story)

def create_doc_summary(entry, output_path):
    """Create a DOC summary of the medical report."""
    doc = docx.Document()

    # Title
    title = doc.add_heading('Medical Report Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary section
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(entry['summary'])

    # Key Information section
    doc.add_heading('Key Information', level=1)
    
    for category, items in entry['key_information'].items():
        # Add category heading
        doc.add_heading(category.capitalize(), level=2)
        
        # Add items
        if items:
            for item in items:
                p = doc.add_paragraph(item)
                p.style = 'List Bullet'
        else:
            doc.add_paragraph(f"No {category} information found")
        
        # Add confidence score
        confidence = entry['confidence_scores'].get(category, 0)
        if confidence:
            p = doc.add_paragraph(f"Confidence: {confidence*100:.1f}%")
            p.style = 'List Bullet'

    # Save document
    doc.save(output_path)

@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/summarize')
def home():
    return render_template('index.html')

@app.route('/history')
def history():
    return render_template('history.html')

@app.route('/api/history')
def get_history():
    """Get history from database."""
    try:
        history = load_history()
        return jsonify(history)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/summarize', methods=['POST'])
def summarize():
    data = request.get_json()
    medical_text = data.get('text', '')
    
    if not medical_text:
        return jsonify({'error': 'No text provided'}), 400
    
    try:
        summary = generate_summary(medical_text)
        key_info, confidence_scores = extract_key_information(medical_text)
        highlighted_text = highlight_medical_terms(medical_text)
        
        # Save to history
        entry = {
            'filename': 'Manual Input',
            'summary': summary,
            'key_info': key_info,
            'confidence_scores': confidence_scores
        }
        save_history(entry)
        
        return jsonify({
            'summary': summary,
            'key_information': key_info,
            'confidence_scores': confidence_scores,
            'highlighted_text': highlighted_text
        })
    except Exception as e:
        return jsonify({'error': f'Error processing text: {str(e)}'}), 500

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        try:
            filename = secure_filename(file.filename)
            file_extension = filename.rsplit('.', 1)[1].lower()
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # Save the file
            file.save(filepath)
            
            # Extract text based on file type
            content = extract_text_from_file(filepath, file_extension)
            
            # Clean up the temporary file
            os.remove(filepath)
            
            # Generate summary and extract key information
            summary = generate_summary(content)
            key_info, confidence_scores = extract_key_information(content)
            highlighted_text = highlight_medical_terms(content)
            
            # Save to history
            entry = {
                'filename': filename,
                'summary': summary,
                'key_info': key_info,
                'confidence_scores': confidence_scores
            }
            save_history(entry)
            
            return jsonify({
                'content': content,
                'summary': summary,
                'key_information': key_info,
                'confidence_scores': confidence_scores,
                'highlighted_text': highlighted_text
            })
            
        except Exception as e:
            # Clean up the file if it exists
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/api/delete_history/<int:entry_id>', methods=['DELETE'])
def delete_history(entry_id):
    """Delete a history entry."""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('DELETE FROM history WHERE id = ?', (entry_id,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/<int:index>')
def export_summary(index):
    try:
        format_type = request.args.get('format', 'pdf').lower()
        history = load_history()
        
        if 0 <= index < len(history):
            entry = history[index]
            
            if format_type == 'pdf':
                # Create PDF
                export_path = os.path.join(app.config['UPLOAD_FOLDER'], f'summary_{index}.pdf')
                create_pdf_summary(entry, export_path)
                return send_file(export_path, as_attachment=True, download_name=f'summary_{index}.pdf')
                
            elif format_type == 'doc':
                # Create DOC
                export_path = os.path.join(app.config['UPLOAD_FOLDER'], f'summary_{index}.docx')
                create_doc_summary(entry, export_path)
                return send_file(export_path, as_attachment=True, download_name=f'summary_{index}.docx')
                
            else:
                # Create CSV (default)
                df = pd.DataFrame({
                    'Category': ['Summary'] + list(entry['key_info'].keys()),
                    'Content': [entry['summary']] + [', '.join(entry['key_info'][k]) for k in entry['key_info'].keys()],
                    'Confidence': ['N/A'] + [f"{entry['confidence_scores'].get(k, 'N/A')*100:.1f}%" for k in entry['key_info'].keys()]
                })
                export_path = os.path.join(app.config['UPLOAD_FOLDER'], f'summary_{index}.csv')
                df.to_csv(export_path, index=False)
                return send_file(export_path, as_attachment=True, download_name=f'summary_{index}.csv')
        
        return jsonify({'error': 'Invalid history index'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/current', methods=['POST'])
def export_current_summary():
    """Export the current summary without requiring a history index."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        format_type = request.args.get('format', 'pdf').lower()
        summary = data.get('summary', '')
        key_information = data.get('key_information', {})
        confidence_scores = data.get('confidence_scores', {})

        entry = {
            'summary': summary,
            'key_information': key_information,
            'confidence_scores': confidence_scores
        }

        if format_type == 'pdf':
            # Create PDF
            export_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_summary.pdf')
            create_pdf_summary(entry, export_path)
            return send_file(export_path, as_attachment=True, download_name='medical_summary.pdf')
            
        elif format_type == 'doc':
            # Create DOC
            export_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_summary.docx')
            create_doc_summary(entry, export_path)
            return send_file(export_path, as_attachment=True, download_name='medical_summary.docx')
            
        else:
            # Create CSV (default)
            df = pd.DataFrame({
                'Category': ['Summary'] + list(key_information.keys()),
                'Content': [summary] + [', '.join(key_information[k]) for k in key_information.keys()],
                'Confidence': ['N/A'] + [f"{confidence_scores.get(k, 'N/A')*100:.1f}%" for k in key_information.keys()]
            })
            export_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_summary.csv')
            df.to_csv(export_path, index=False)
            return send_file(export_path, as_attachment=True, download_name='medical_summary.csv')

    except Exception as e:
        return jsonify({'error': f'Error exporting summary: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True) 